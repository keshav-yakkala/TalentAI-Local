"""
Resume Parser Service (Phase 3)
Wraps PDF/DOCX/TXT extraction in a proper service class.
Upgrade from PyPDF2 → PyMuPDF (primary) with fallback chain.
"""
from __future__ import annotations

import io
import re
import unicodedata
from pathlib import Path

from app.core.exceptions import ResumeParsingError, UnsupportedFileError
from app.core.logging import get_logger

logger = get_logger(__name__)

SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class ResumeParserService:
    """
    Service for extracting raw text from resume files.
    
    Extraction chain:
      PDF  → PyMuPDF (fitz) → pdfplumber → PyPDF2
      DOCX → python-docx
      TXT  → utf-8 with error replacement
    """

    async def parse(
        self,
        file_content: bytes,
        filename: str,
        mime_type: str,
    ) -> tuple[str, float]:
        """
        Parse resume bytes and return (raw_text, confidence).
        confidence is 0.0–1.0 based on text quality.
        
        Raises:
            UnsupportedFileError: file type not supported
            ResumeParsingError: all parsers failed
        """
        ext = Path(filename).suffix.lower()

        if mime_type not in SUPPORTED_MIME_TYPES and ext not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileError(
                f"File type '{mime_type}' / '{ext}' is not supported",
                {"filename": filename, "mime_type": mime_type},
            )

        logger.info("Parsing resume", filename=filename, mime_type=mime_type, size_bytes=len(file_content))

        if mime_type == "application/pdf" or ext == ".pdf":
            text, confidence = await self._parse_pdf(file_content, filename)
        elif mime_type.endswith("wordprocessingml.document") or ext == ".docx":
            text, confidence = self._parse_docx(file_content, filename)
        else:
            text, confidence = self._parse_txt(file_content, filename)

        text = self._clean_text(text)

        if len(text.strip()) < 50:
            raise ResumeParsingError(
                "Extracted text is too short — may be a scanned PDF requiring OCR",
                {"filename": filename, "text_length": len(text)},
            )

        logger.info("Parse completed", filename=filename, text_length=len(text), confidence=confidence)
        return text, confidence

    async def _parse_pdf(self, content: bytes, filename: str) -> tuple[str, float]:
        """PDF extraction with PyMuPDF → pdfplumber → PyPDF2 fallback chain."""
        # Attempt 1: PyMuPDF (best quality)
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            pages_text = []
            for page in doc:
                pages_text.append(page.get_text())
            doc.close()
            text = "\n".join(pages_text)
            if len(text.strip()) >= 100:
                logger.debug("PDF parsed with PyMuPDF", filename=filename)
                return text, 0.95
        except ImportError:
            logger.warning("PyMuPDF not installed, trying pdfplumber")
        except Exception as exc:
            logger.warning("PyMuPDF failed", filename=filename, error=str(exc))

        # Attempt 2: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages_text = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages_text)
            if len(text.strip()) >= 100:
                logger.debug("PDF parsed with pdfplumber", filename=filename)
                return text, 0.85
        except ImportError:
            logger.warning("pdfplumber not installed, trying PyPDF2")
        except Exception as exc:
            logger.warning("pdfplumber failed", filename=filename, error=str(exc))

        # Attempt 3: PyPDF2 (legacy fallback)
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            pages_text = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages_text)
            if len(text.strip()) >= 50:
                logger.debug("PDF parsed with PyPDF2", filename=filename)
                return text, 0.70
        except Exception as exc:
            logger.error("All PDF parsers failed", filename=filename, error=str(exc))

        raise ResumeParsingError(
            "All PDF parsers failed. File may be scanned or corrupted.",
            {"filename": filename},
        )

    def _parse_docx(self, content: bytes, filename: str) -> tuple[str, float]:
        """DOCX extraction with python-docx."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

            text = "\n".join(paragraphs)
            logger.debug("DOCX parsed", filename=filename, paragraphs=len(paragraphs))
            return text, 0.90
        except Exception as exc:
            raise ResumeParsingError(
                f"DOCX parsing failed: {exc}",
                {"filename": filename},
            ) from exc

    def _parse_txt(self, content: bytes, filename: str) -> tuple[str, float]:
        """Plain text extraction with encoding detection."""
        for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                text = content.decode(encoding)
                return text, 0.80
            except (UnicodeDecodeError, LookupError):
                continue
        # Last resort: decode with error replacement
        text = content.decode("utf-8", errors="replace")
        return text, 0.60

    def _clean_text(self, text: str) -> str:
        """Normalize text: whitespace, encoding artifacts, control characters."""
        if not text:
            return ""
        # Normalize unicode (e.g., ligatures, special chars)
        text = unicodedata.normalize("NFKC", text)
        # Remove null bytes and control characters (keep newlines and tabs)
        text = "".join(ch for ch in text if ch.isprintable() or ch in "\n\t")
        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Collapse excessive blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Collapse excessive spaces
        text = re.sub(r"[ \t]+", " ", text)
        # Strip leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)
        return text.strip()
